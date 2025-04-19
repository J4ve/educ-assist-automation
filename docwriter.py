from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re
from docx.oxml import OxmlElement

FILENAME_TXT = "educational_assistance_list.txt"
FILENAME_DOCX = "educational_assistance_list.docx"

def load_entries(filename):
    if not os.path.exists(filename):
        return []
    with open(filename, "r", encoding="utf-8") as f:
        content = f.read()
    entries = []
    blocks = re.findall(r"==== Entry #\d+ ====\n(.*?)\n--------------------------", content, re.DOTALL)
    for block in blocks:
        parent = {}
        student = {}
        p_block = re.search(r"Parent/Guardian Information:\n(.*?)\n\nStudent Information:", block, re.DOTALL)
        s_block = re.search(r"Student Information:\n(.*)", block, re.DOTALL)
        if p_block:
            for line in p_block.group(1).strip().splitlines():
                if ':' in line:
                    k, v = line.strip().split(":", 1)
                    parent[k.strip()] = v.strip()
        if s_block:
            for line in s_block.group(1).strip().splitlines():
                if ':' in line:
                    k, v = line.strip().split(":", 1)
                    student[k.strip()] = v.strip()
        entries.append((parent, student))
    return entries

def format_entry_text(parent, student):
    lines = ["Parent/Guardian Information:"]
    for k, v in parent.items():
        lines.append(f"{k}: {v}")
    lines.append("")
    lines.append("Student Information:")
    for k, v in student.items():
        lines.append(f"{k}: {v}")
    return "\n".join(lines)

def write_docx(entries, filename=FILENAME_DOCX):
    doc = Document()

    # Set the paper size to Legal (8.5 x 14 inches)
    section = doc.sections[0]
    section.page_width = Inches(8.5)  # 8.5 inches wide
    section.page_height = Inches(14)  # 14 inches tall

    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    for i in range(0, len(entries), 4):  # Every 4 entries
        table = doc.add_table(rows=2, cols=2)  # 2 rows, 2 columns for 4 instances per page
        table.autofit = True

        # Fill in the table with data
        for row in range(2):  # 2 rows
            for col in range(2):  # 2 columns
                idx = i + (row * 2) + col
                if idx < len(entries):
                    parent, student = entries[idx]
                    text = format_entry_text(parent, student)
                    cell = table.cell(row, col)
                    para = cell.paragraphs[0]
                    run = para.add_run(text)
                    run.font.size = Pt(10)
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        if i + 4 < len(entries):
            doc.add_page_break()  # Force a new page after 4 entries

    doc.save(filename)
    print(f"Saved DOCX with {len(entries)} entries: {filename}")

if __name__ == "__main__":
    entries = load_entries(FILENAME_TXT)
    if not entries:
        print("No entries found. Run the input script first.")
    else:
        write_docx(entries)
